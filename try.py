from docx import Document
from bs4 import BeautifulSoup

# HTML content (you can replace this with reading from a file)
html_content = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <title>Ashish Singh - Resume</title>
</head>
<body>

<div class="resume">
    <div class="top">
        <div class="photo">
            <!-- <img src="your-photo.jpg" alt="Profile Photo" /> -->
        </div>
        <div class="top-info">
            <h1>Ashish Singh</h1>
            <div class="line"></div>
            <h3>Full Stack Developer</h3>
        </div>
    </div>

    <div class="content">
        <div class="left-col">
            <section>
                <h2>CONTACT</h2>
                <p>📞 ‪‪+91 7208523099‬‬</p>
                <p>📧 Ashishroysingh135@gmail.com</p>
                <p>🌍 Mumbai, India</p>
                <p>💻 ‪github.com/SinghAshish87‬</p>
            </section>

            <section>
                <h2>ABOUT ME</h2>
                <p>BSc Computer Science graduate with a passion for technology and a strong interest in developing practical solutions. Eager to begin a professional journey in the IT field and grow through continuous learning and real-world experience.</p>
            </section>

            <section>
                <h2>SKILLS</h2>
                <ul>
                    <li>HTML / CSS / JavaScript</li>
                    <li>Python</li>
                    <li>SQL / MongoDB</li>
                    <li>MS Word</li>
                    <li>PowerPoint</li>
                    <li>Java</li>
                    <li>React.js</li>
                    <li>Next.js</li>
                </ul>
            </section>

            <section>
                <h2>LANGUAGES</h2>
                <p>English – Fluent</p>
                <p>Hindi – Native</p>
            </section>
        </div>

        <div class="right-col">
            <section>
                <h2>EDUCATION</h2>
                <p><strong>B.Sc Computer Science</strong><br />GPM Degree College — 2025</p>
                <p><strong>HSC (Science)</strong><br />RD National College — 2022</p>
                <p><strong>SSC</strong><br />St. Anne’s High School — 2020</p>
            </section>

            <section>
                <h2>PROJECTS</h2>
                <p><strong>Portfolio Website</strong><br />Designed and developed a full responsive personal portfolio website.</p>
                <p><strong>Industrial Project – YMCA</strong><br />Participated in analysis & UI / backend development.</p>
                <p><strong>Sports Shop Website</strong><br />Product listing + UI-based Functional pages.</p>
                <p><strong>Technologies Used:</strong> HTML, CSS, JavaScript, MongoDB</p>
            </section>

            <section>
                <h2>EXPERIENCE</h2>
                <p><strong>Fresher</strong><br />Open to internships & full-time roles.</p>
            </section>
        </div>
    </div>
</div>

</body>
</html>'''

# Parse HTML with BeautifulSoup
soup = BeautifulSoup(html_content, 'html.parser')

# Create a new Word Document
doc = Document()
doc.add_heading('Ashish Singh - Resume', 0)

# Extract sections and add to the Word document
def add_section_to_doc(section_title, content):
    doc.add_heading(section_title, level=1)
    if isinstance(content, str):
        doc.add_paragraph(content)
    else:
        for item in content:
            doc.add_paragraph(item, style='List Bullet')

# Extract and add contact info
contact_section = soup.find('section', text="CONTACT")
contact_info = [item.get_text() for item in contact_section.find_all('p')]
add_section_to_doc('CONTACT', contact_info)

# Extract and add about me
about_me_section = soup.find('section', text="ABOUT ME")
about_me_text = about_me_section.find('p').get_text(strip=True)
add_section_to_doc('ABOUT ME', about_me_text)

# Extract and add skills
skills_section = soup.find('section', text="SKILLS")
skills = [item.get_text() for item in skills_section.find_all('li')]
add_section_to_doc('SKILLS', skills)

# Extract and add languages
languages_section = soup.find('section', text="LANGUAGES")
languages = [item.get_text() for item in languages_section.find_all('p')]
add_section_to_doc('LANGUAGES', languages)

# Extract and add education
education_section = soup.find('section', text="EDUCATION")
education = [item.get_text(strip=True) for item in education_section.find_all('p')]
add_section_to_doc('EDUCATION', education)

# Extract and add projects
projects_section = soup.find('section', text="PROJECTS")
projects = [item.get_text(strip=True) for item in projects_section.find_all('p')]
add_section_to_doc('PROJECTS', projects)

# Extract and add experience
experience_section = soup.find('section', text="EXPERIENCE")
experience_text = experience_section.find('p').get_text(strip=True)
add_section_to_doc('EXPERIENCE', experience_text)

# Save the document as a .docx file
doc.save('Ashish_Singh_Resume.docx')

print("Word document saved successfully!")
