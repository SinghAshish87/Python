from docx import Document
from bs4 import BeautifulSoup

# Parse HTML with BeautifulSoup
soup = BeautifulSoup(html_content, 'html.parser')

doc = Document()
doc.add_heading('Ashish Singh - Resume', 0)

def get_section(title):
    h2 = soup.find('h2', string=title)
    if h2:
        return h2.parent
    return None

def add_section_to_doc(section_title, content):
    doc.add_heading(section_title, level=1)
    if isinstance(content, str):
        doc.add_paragraph(content)
    else:
        for item in content:
            doc.add_paragraph(item, style='List Bullet')

# CONTACT
contact_sec = get_section("CONTACT")
contact_info = [p.get_text() for p in contact_sec.find_all('p')]
add_section_to_doc("CONTACT", contact_info)

# ABOUT ME
about_sec = get_section("ABOUT ME")
about_text = about_sec.find('p').get_text()
add_section_to_doc("ABOUT ME", about_text)

# SKILLS
skills_sec = get_section("SKILLS")
skills = [li.get_text() for li in skills_sec.find_all('li')]
add_section_to_doc("SKILLS", skills)

# LANGUAGES
lang_sec = get_section("LANGUAGES")
languages = [p.get_text() for p in lang_sec.find_all('p')]
add_section_to_doc("LANGUAGES", languages)

# EDUCATION
edu_sec = get_section("EDUCATION")
education = [p.get_text() for p in edu_sec.find_all('p')]
add_section_to_doc("EDUCATION", education)

# PROJECTS
proj_sec = get_section("PROJECTS")
projects = [p.get_text() for p in proj_sec.find_all('p')]
add_section_to_doc("PROJECTS", projects)

# EXPERIENCE
exp_sec = get_section("EXPERIENCE")
experience = exp_sec.find('p').get_text()
add_section_to_doc("EXPERIENCE", experience)

doc.save("Ashish_Singh_Resume.docx")
print("Word document saved successfully!")


